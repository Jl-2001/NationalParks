import requests
import docx
import random
from pprint import pprint

park_doc = docx.Document()
park_doc.add_paragraph('national park', 'Heading 1')
#this creates the heading which is called national park.

parklist_url = 'https://national-parks-1150.azurewebsites.net/api/list'

parklist_response = requests.get(parklist_url).json()

# print(parklist_response)
#notes from video, use descriptive variable names to help keep things organized

five_random_parks = random.sample(parklist_response, 5)
# print(five_random_parks)

for park in five_random_parks:
    # print(park) #this is the dictionary with the park name and park code

    #example url for park details is https://national-parks-1150.azurewebsites.net/api/ACAD
    #remember the urls are strings.
    base_park_detail_url = 'https://national-parks-1150.azurewebsites.net/api/'
    park_code = park['park_code']
    park_name = park['name']
    park_doc.add_paragraph(park_name)
    #here we created a doc that requests five parks and the prints out the name into a word doc

    print(base_park_detail_url)
    print(park_code)
    park_detail_url = base_park_detail_url + park_code
    print(park_detail_url)

    park_detail_response = requests.get(park_detail_url).json()
    pprint(park_detail_response)


park_doc.save('parkguide.docx')