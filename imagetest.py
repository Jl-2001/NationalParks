import random
import docx
import requests
from pprint import pprint

from docx.shared import Inches

from worddocument import description, activities, location

national_park_doc = docx.Document()
national_park_doc.add_paragraph('National Park Guide', 'Heading 1')

national_park_url = 'https://national-parks-1150.azurewebsites.net/api/list'
park_response = requests.get(national_park_url).json()

# pprint(park_response)

five_random_national = random.sample(park_response, 5)
# pprint(five_random_national)

for park in five_random_national:
    #pprint(park) #for example {'name': 'Yosemite National Park', 'park_code': 'YOSE'}

    park_code = park['park_code']
    base_park_url = 'https://national-parks-1150.azurewebsites.net/api/' + park_code

    # parks_urls = base_park_url + park_code this is an extra, not needed. caused error
    print(base_park_url)
    parkrequests_api = requests.get(base_park_url)
    park_data = parkrequests_api.json()

    park_name = park_data['name']
    #adding park as a heading name
    park_activities = park_data['activities']
    #pinging the api to access the images.
    park_images = park_data['nps_park_images']

    national_park_doc.add_paragraph(park_name, 'Heading 2')

    # pprint(parks_urls)
    for images in park_images:
        image_title = images['title']
        image_credit = images['credit']
        image_url = images['url']
        nps_images_list = list(park_images)
        image_response = requests.get(image_url)
        url_piece = image_url.split('/')
        file_name = url_piece.pop()

        if nps_images_list == '':
            national_park_doc.add_paragraph('No image found', 'Caption')
            #checks to see if there is no images
        for index, url in enumerate(nps_images_list):#enumerate: makes the list of images in numbers
            with open(file_name, 'wb')as file:#write binary
                for chunk in image_response.iter_content():#breaks images in chunks
                    file.write(chunk)
        national_park_doc.add_picture(file_name, width=Inches(4), height=Inches(4))


        # pprint(park_data) this prints the parks descriptions but not precise enough
        #below will add the description if it available.

    describe_park = park_data.get('description', 'No description available')
    national_park_doc.add_paragraph(f'description ${describe_park}')


    #adding activities if they available
    activities = park_data.get('activities', [])
    if activities:
        activities_list = ', '.join(activities)
        national_park_doc.add_paragraph(f'activities are: ${activities_list}')
    else:
        national_park_doc.add_paragraph('There are no activities available')

    #add additional information
    location = park_data('location', 'location no provided')
    national_park_doc.add_paragraph(f'location: ${location}')


    image_list = park_data.get('name', 'no image available')
    national_park_doc.add_picture(image_list)
    # national_park_doc.add_paragraph(park_data)

national_park_doc.save('NationalParkDocTest.docx')
