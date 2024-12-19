import docx

# document = docx.Document()
#
# document.add_paragraph('hello world', 'Title')
#
#
# document.add_paragraph('By Jorge Lazaro', 'Heading 1')
#
# document.add_paragraph('this is a word doc for my python class')
#
# document.save('helloworld.docx')


import random
import docx
import requests

# Initialize Word document
national_park_doc = docx.Document()
national_park_doc.add_paragraph('National Park Guide', 'Title')

# Fetch list of parks
national_park_url = 'https://national-parks-1150.azurewebsites.net/api/list'
park_response = requests.get(national_park_url).json()
#json for the dictionary from the url in json format

# Randomly select five parks
five_random_national = random.sample(park_response, 5)

# Fetch details for each selected park and add to Word document
for park in five_random_national:
    base_park_url = 'https://national-parks-1150.azurewebsites.net/api/'
    park_code = park['park_code']
    park_name = park['name']

    # Add park name as a heading
    national_park_doc.add_heading(park_name, level=2)

    # Fetch park details
    parks_urls = base_park_url + park_code
    parkapi_response = requests.get(parks_urls)
    #PROVIDES ID THE PARK API DOES RESPOND AND IS AT LEAST A 200 THEN THE DATA DISPLAYS IN JSON FORMAT
    if parkapi_response.status_code == 200:
        park_data = parkapi_response.json()

        # Add description if available
        #WE DO THIS BY CREATING DESCRIPTION CAPTION AND IF NONE IS AVAILABLE THEN IT
        #WILL DISPLAY THAT NONE ARE AVAILABLE
        description = park_data.get('description', 'No description available.')
        national_park_doc.add_paragraph(f"Description: {description}")

        # Add activities if available
        #WE DO THIS BY CREATING AN IF STATEMENT AND IF THE ACTIVITY DOES NOT EXIST THEN IT WILL
        #DISPLAY THAT NONE ARE AVAILABLE
        activities = park_data.get('activities', [])
        if activities:
            activities_list = ", ".join(activities)
            national_park_doc.add_paragraph(f"Activities: {activities_list}")
        else:
            national_park_doc.add_paragraph("Activities: None available.")

        # Add additional information as needed
        location = park_data.get('location', 'Location not provided.')
        national_park_doc.add_paragraph(f"Location: {location}")

        national_park_doc.add_paragraph('contact info')
        contact_infodict = park_data['contact_info']

        print(contact_infodict)



        # for key,value in contact_infodict.items():
        #     national_park_doc.add_paragraph(key, 'Heading 3')
        #     national_park_doc.add_paragraph(value)

        #EXTRACTING DATA FROM THE API TO SHOW THE ADDRESS AND DISPLAYING THAT URL OF THE ADDRESS
        address_text = contact_infodict['address']
        park_website_text = contact_infodict['url']

        #PROVIDING A CAPTION FOR THE ADDRESS AS WELL AS THE CORRECT ADDRESS DISPLAYING BENEATH THE CAPTION
        national_park_doc.add_paragraph('Address', 'Heading 3')
        national_park_doc.add_paragraph(address_text)

        # national_park_doc.add_paragraph('Website', 'Heading 3')
        # national_park_doc.add_paragraph(park_website_text)

        # THIS IS TO EXTRACT THE DATA FROM THE API THAT HAS THE IMAGES
        park_image_test = park_data['nps_park_images']
        # for park_image in park_image_test:
        #     print(park_image)
        for one_park_at_a_time in park_image_test:
            image_caption = one_park_at_a_time['caption']
            national_park_doc.add_paragraph(image_caption)

            image_for_park_url = one_park_at_a_time['url']

            print(image_for_park_url)

            national_park_doc.add_paragraph('Image', 'Heading 3')#heading caption for each image
            image_response = requests.get(image_for_park_url)
            with open('park_image.jpg', 'wb') as file:
                for chunk in image_response.iter_content():
                    #now we will overwrite park images for each images.
                    file.write(chunk)


            national_park_doc.add_picture('park_image.jpg', width=docx.shared.Inches(6))







# Save the Word document
national_park_doc.save('testingparksonetwo.docx')